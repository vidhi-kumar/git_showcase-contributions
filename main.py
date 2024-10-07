import os
import io
import json
from openai import OpenAI
import streamlit as st
from docx import Document
from get_repo_list import get_repo_list
from get_user_contributions import get_user_contributions
from get_final_prompt import get_final_prompt
from get_final_results import get_final_results

def create_word_doc(contributions, github_username):
    doc = Document()
    doc.add_heading(f'GitHub Contributions for {github_username}', level=0)

    if isinstance(contributions, list):
        for contribution in contributions:
            if 'repository_name' in contribution:
                doc.add_heading(contribution['repository_name'], level=1)

            if 'contribution_summary' in contribution:
                for summary in contribution['contribution_summary']:
                    doc.add_paragraph(summary, style='List Bullet')
    else:
        doc.add_paragraph("No valid contributions found.")
    
    return doc



def main():

    # st.title("git showcase-contributions")
    st.markdown("<h1 style='color:#003366;'>git showcase-contributions</h1>", unsafe_allow_html=True)

    # instructions
    with st.expander("Instructions to create GitHub Fine-grained Personal Access Token"):
        st.markdown("""
        1. Go to your [GitHub Personal Access Token settings](https://github.com/settings/tokens?type=beta).
        2. Click on **Generate new token** and choose **Fine-grained token**.
        3. Select a custom expiration date (example: 7 days)
        4. Set the **Repository access** to **Public Repositories (read-only)**.
        5. Generate the token and save it securely. Use the token here.
                    
        **Important:** The token is configured to grant access to **public repositories only** to ensure compliance with GitHub’s access control policies and to avoid unauthorized or restricted access to private repositories.
        
        **Data Usage Notice:** This application will access and utilize information from your pull requests, issues, and the README files of your repositories. This data will be used to generate contribution summaries and related insights.
        """)

    openai_api_key = os.getenv('OPENAPI_KEY')
    client = OpenAI(api_key=openai_api_key)


    col1, _ = st.columns([3.5, 2])
    with col1:
        github_username = st.text_input("github username", "", key="username_input")
        github_access_token = st.text_input("github access token", type="password", key="access_token")

    
    col1, _ = st.columns([1.5, 2])
    with col1:
        expertise = ["Data Science", "Machine Learning", "Artificial Intelligence", "Frontend Development", "Backend Development", "Fullstack Development", "Devops Engineering", "Cloud Engineering", "UI/UX"]  # List of models
        expertise = st.selectbox("Select your expertise", expertise, index=expertise.index("Fullstack Development"))
        st.caption(f"Selected Expertise: {expertise}")
    
    col1, _ = st.columns([1, 2])
    with col1:
        summary_size = ["3", "5", "8"]  
        summary_size = st.selectbox("Summary sentences per repository", summary_size, index=summary_size.index("3"))
        st.caption(f"Selected Summary Length: {summary_size}")

    col1, _ = st.columns([1.5, 2])
    with col1:
        models = ["gpt-4o-mini", "gpt-3.5-turbo"]  
        selected_model = st.selectbox("Select Model", models, index=models.index("gpt-4o-mini"))
        st.caption(f"Selected Model: {selected_model}")

    st.write("######")
    if st.button("Showcase My Contributions"):
        if not github_username:
            st.warning("Please enter your GitHub username.")
        elif not github_access_token:
            st.warning("Please enter your GitHub access token.")
        elif not summary_size:
            st.warning("Please enter Summary sentences per repository")
        elif not expertise:
            st.warning("Please enter your expertise")
        else:
            percent_complete = 0
            progress_text = "Operation in progress. Please wait."
            my_bar = st.progress(percent_complete, text=progress_text)
            my_bar.progress(percent_complete, text="listing down all the repositories you are a part of...")
            repo_list = get_repo_list(access_token=github_access_token)
            repo_list = []
        
            if len(repo_list) == 0:
                my_bar.empty()
                st.warning("You don't have any contributions yet..")
            else:
                print("repo_list")
                print(repo_list)
                
                percent_complete += 25
                my_bar.progress(percent_complete, text="scanning through your contributions...")
                user_contributions = get_user_contributions(repository_list=repo_list, username=github_username, access_token=github_access_token)
                print("user_contributions")
                print(user_contributions)
                
                percent_complete += 25
                my_bar.progress(percent_complete , text="generating prompt...")
                final_prompt = get_final_prompt(user_contributions=user_contributions, summary_size=summary_size, expertise=expertise)
                print("final_prompt")
                print(final_prompt)

                percent_complete += 25
                my_bar.progress(percent_complete, text="feeding prompt to model...")
                final_results = get_final_results(final_prompt=final_prompt, client=client, model=selected_model)
                contributions = json.loads(final_results)
                print("contributions")
                print(contributions)

                percent_complete += 25
                my_bar.progress(percent_complete, text="generating results...")
                my_bar.empty()


                # Create the Word document
                doc = create_word_doc(contributions, github_username)
                
                # Save the document to an in-memory BytesIO stream
                word_io = io.BytesIO()  # Create BytesIO object
                doc.save(word_io)  # Save document to the BytesIO stream
                word_io.seek(0)  # Go to the start of the stream

                # Provide the download button in Streamlit
                st.download_button(
                    label="Download GitHub Contributions",
                    data=word_io,
                    file_name=f"{github_username}_contributions.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                # Display contributions
                st.title(f"Contributions by {github_username}")

                for contribution in contributions:
                    with st.expander(contribution["repository_name"]):
                        for summary in contribution["contribution_summary"]:
                            st.markdown(f"- {summary}")

        

if __name__ == "__main__":
    main()